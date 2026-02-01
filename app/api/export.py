"""
IC Memo Export API for DealFlow AI Copilot.

Export analysis results to PDF and DOCX formats.
"""

import io
import json
from datetime import datetime
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from app.utils.logger import logger

router = APIRouter(prefix="/export", tags=["export"])

# In-memory results store (shared with deals.py)
# In production, this would read from the database
from app.api.deals import _analysis_results


@router.get(
    "/{analysis_id}/pdf",
    summary="Export IC Memo as PDF",
    description="Generate a professional PDF of the IC memo",
)
async def export_pdf(analysis_id: str) -> StreamingResponse:
    """Export an IC memo as a formatted PDF document."""
    memo = _analysis_results.get(analysis_id)
    if not memo:
        raise HTTPException(status_code=404, detail=f"Analysis not found: {analysis_id}")

    try:
        pdf_buffer = _generate_pdf(memo.model_dump())

        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="IC_Memo_{memo.company_name}_{analysis_id}.pdf"'
            },
        )
    except Exception as e:
        logger.error(f"PDF export failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.get(
    "/{analysis_id}/docx",
    summary="Export IC Memo as DOCX",
    description="Generate a Word document of the IC memo",
)
async def export_docx(analysis_id: str) -> StreamingResponse:
    """Export an IC memo as a Word document."""
    memo = _analysis_results.get(analysis_id)
    if not memo:
        raise HTTPException(status_code=404, detail=f"Analysis not found: {analysis_id}")

    try:
        docx_buffer = _generate_docx(memo.model_dump())

        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="IC_Memo_{memo.company_name}_{analysis_id}.docx"'
            },
        )
    except Exception as e:
        logger.error(f"DOCX export failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


@router.get(
    "/{analysis_id}/json",
    summary="Export IC Memo as JSON",
    description="Download the raw JSON analysis data",
)
async def export_json(analysis_id: str) -> StreamingResponse:
    """Export raw IC memo data as JSON."""
    memo = _analysis_results.get(analysis_id)
    if not memo:
        raise HTTPException(status_code=404, detail=f"Analysis not found: {analysis_id}")

    json_str = memo.model_dump_json(indent=2)
    buffer = io.BytesIO(json_str.encode("utf-8"))

    return StreamingResponse(
        buffer,
        media_type="application/json",
        headers={
            "Content-Disposition": f'attachment; filename="IC_Memo_{memo.company_name}_{analysis_id}.json"'
        },
    )


def _generate_pdf(memo_data: dict[str, Any]) -> io.BytesIO:
    """Generate a professional PDF from IC memo data."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch, mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=30 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "ICTitle",
        parent=styles["Title"],
        fontSize=22,
        spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e"),
    )
    heading_style = ParagraphStyle(
        "ICHeading",
        parent=styles["Heading1"],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor("#16213e"),
        borderWidth=1,
        borderColor=colors.HexColor("#0f3460"),
        borderPadding=4,
    )
    subheading_style = ParagraphStyle(
        "ICSubheading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=4,
        textColor=colors.HexColor("#0f3460"),
    )
    body_style = ParagraphStyle(
        "ICBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        "ICBullet",
        parent=body_style,
        leftIndent=20,
        bulletIndent=10,
    )

    elements = []

    # Title
    company = memo_data.get("company_name", "Unknown Company")
    elements.append(Paragraph(f"Investment Committee Memo", title_style))
    elements.append(Paragraph(f"<b>{company}</b>", styles["Heading2"]))
    elements.append(Paragraph(
        f"Prepared by: DealFlow AI Copilot | Date: {datetime.now().strftime('%B %d, %Y')}",
        body_style,
    ))
    elements.append(Spacer(1, 12))

    # Executive Summary
    exec_summary = memo_data.get("executive_summary", {})
    if exec_summary:
        elements.append(Paragraph("Executive Summary", heading_style))
        if exec_summary.get("company_overview"):
            elements.append(Paragraph(exec_summary["company_overview"], body_style))

        rec = exec_summary.get("recommendation", "N/A")
        elements.append(Paragraph(f"<b>Recommendation:</b> {rec}", body_style))

        if exec_summary.get("recommendation_rationale"):
            elements.append(Paragraph(exec_summary["recommendation_rationale"], body_style))

        highlights = exec_summary.get("investment_highlights", [])
        if highlights:
            elements.append(Paragraph("Investment Highlights:", subheading_style))
            for h in highlights:
                elements.append(Paragraph(f"• {h}", bullet_style))

        concerns = exec_summary.get("key_concerns", [])
        if concerns:
            elements.append(Paragraph("Key Concerns:", subheading_style))
            for c in concerns:
                elements.append(Paragraph(f"• {c}", bullet_style))

    # Extraction Summary
    extraction = memo_data.get("extraction_result", {})
    if extraction:
        elements.append(Paragraph("Company Overview", heading_style))

        info_data = [
            ["Field", "Value"],
            ["Industry", extraction.get("industry", "N/A")],
            ["Stage", extraction.get("stage", "N/A")],
            ["Business Model", extraction.get("business_model", "N/A")],
            ["Founded", str(extraction.get("founded_year", "N/A"))],
            ["Headquarters", extraction.get("headquarters", "N/A")],
        ]
        info_table = Table(info_data, colWidths=[2 * inch, 4 * inch])
        info_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16213e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f0f0")]),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 12))

    # Financials
    financials = extraction.get("financials", {})
    if financials:
        elements.append(Paragraph("Financial Metrics", heading_style))

        fin_rows = [["Metric", "Value"]]
        if financials.get("revenue"):
            rev = financials["revenue"]
            fin_rows.append(["Revenue", f"${rev.get('amount', 0)}{rev.get('unit', '')}"])
        if financials.get("revenue_growth_rate"):
            fin_rows.append(["Revenue Growth", f"{financials['revenue_growth_rate']:.0%}"])
        if financials.get("gross_margin"):
            fin_rows.append(["Gross Margin", f"{financials['gross_margin']:.0%}"])
        if financials.get("monthly_burn_rate"):
            burn = financials["monthly_burn_rate"]
            fin_rows.append(["Monthly Burn", f"${burn.get('amount', 0)}{burn.get('unit', '')}"])
        if financials.get("runway_months"):
            fin_rows.append(["Runway", f"{financials['runway_months']} months"])

        if len(fin_rows) > 1:
            fin_table = Table(fin_rows, colWidths=[2.5 * inch, 3.5 * inch])
            fin_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f3460")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f0f0")]),
            ]))
            elements.append(fin_table)

    # Risk Assessment
    risk_result = memo_data.get("risk_result", {})
    if risk_result:
        elements.append(Paragraph("Risk Assessment", heading_style))
        elements.append(Paragraph(
            f"<b>Overall Risk Score:</b> {risk_result.get('overall_risk_score', 0):.1f}/10",
            body_style,
        ))
        elements.append(Paragraph(
            f"<b>Recommendation:</b> {risk_result.get('risk_adjusted_recommendation', 'N/A')}",
            body_style,
        ))

        deal_breakers = risk_result.get("deal_breakers", [])
        if deal_breakers:
            elements.append(Paragraph("Deal Breakers:", subheading_style))
            for db in deal_breakers:
                elements.append(Paragraph(f"⚠ {db}", bullet_style))

    # Valuation
    val_result = memo_data.get("valuation_result", {})
    if val_result:
        elements.append(Paragraph("Valuation", heading_style))

        low = val_result.get("valuation_range_low", {})
        mid = val_result.get("valuation_range_mid", {})
        high = val_result.get("valuation_range_high", {})

        elements.append(Paragraph(
            f"<b>Valuation Range:</b> "
            f"${low.get('amount', 0):.1f}M - ${mid.get('amount', 0):.1f}M - ${high.get('amount', 0):.1f}M",
            body_style,
        ))

        if val_result.get("ask_vs_valuation"):
            elements.append(Paragraph(
                f"<b>vs. Ask:</b> {val_result['ask_vs_valuation']}",
                body_style,
            ))

    # Final Recommendation
    elements.append(Paragraph("Final Recommendation", heading_style))
    elements.append(Paragraph(
        f"<b>{memo_data.get('final_recommendation', 'N/A')}</b> "
        f"(Conviction: {memo_data.get('conviction_level', 'N/A')})",
        body_style,
    ))

    # Next Steps
    next_steps = exec_summary.get("next_steps", [])
    if next_steps:
        elements.append(Paragraph("Next Steps:", subheading_style))
        for i, step in enumerate(next_steps, 1):
            elements.append(Paragraph(f"{i}. {step}", bullet_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer


def _generate_docx(memo_data: dict[str, Any]) -> io.BytesIO:
    """Generate a Word document from IC memo data."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Title
    company = memo_data.get("company_name", "Unknown Company")
    title = doc.add_heading("Investment Committee Memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(company)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: DealFlow AI Copilot | Date: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_paragraph()  # Spacer

    # Executive Summary
    exec_summary = memo_data.get("executive_summary", {})
    if exec_summary:
        doc.add_heading("Executive Summary", level=1)
        if exec_summary.get("company_overview"):
            doc.add_paragraph(exec_summary["company_overview"])

        rec = exec_summary.get("recommendation", "N/A")
        p = doc.add_paragraph()
        p.add_run("Recommendation: ").bold = True
        p.add_run(str(rec))

        if exec_summary.get("recommendation_rationale"):
            doc.add_paragraph(exec_summary["recommendation_rationale"])

        highlights = exec_summary.get("investment_highlights", [])
        if highlights:
            doc.add_heading("Investment Highlights", level=2)
            for h in highlights:
                doc.add_paragraph(h, style="List Bullet")

        concerns = exec_summary.get("key_concerns", [])
        if concerns:
            doc.add_heading("Key Concerns", level=2)
            for c in concerns:
                doc.add_paragraph(c, style="List Bullet")

    # Company Overview
    extraction = memo_data.get("extraction_result", {})
    if extraction:
        doc.add_heading("Company Overview", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"

        fields = [
            ("Industry", extraction.get("industry", "N/A")),
            ("Stage", extraction.get("stage", "N/A")),
            ("Business Model", extraction.get("business_model", "N/A")),
            ("Founded", str(extraction.get("founded_year", "N/A"))),
            ("HQ", extraction.get("headquarters", "N/A")),
        ]
        for field, value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)

    # Financials
    financials = extraction.get("financials", {})
    if financials:
        doc.add_heading("Financial Metrics", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"

        if financials.get("revenue"):
            rev = financials["revenue"]
            row = table.add_row().cells
            row[0].text = "Revenue"
            row[1].text = f"${rev.get('amount', 0)}{rev.get('unit', '')}"
        if financials.get("revenue_growth_rate") is not None:
            row = table.add_row().cells
            row[0].text = "Growth Rate"
            row[1].text = f"{financials['revenue_growth_rate']:.0%}"
        if financials.get("gross_margin") is not None:
            row = table.add_row().cells
            row[0].text = "Gross Margin"
            row[1].text = f"{financials['gross_margin']:.0%}"

    # Risk Assessment
    risk = memo_data.get("risk_result", {})
    if risk:
        doc.add_heading("Risk Assessment", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Overall Risk Score: {risk.get('overall_risk_score', 0):.1f}/10").bold = True
        doc.add_paragraph(f"Recommendation: {risk.get('risk_adjusted_recommendation', 'N/A')}")

        for db in risk.get("deal_breakers", []):
            doc.add_paragraph(f"⚠ {db}", style="List Bullet")

    # Valuation
    val = memo_data.get("valuation_result", {})
    if val:
        doc.add_heading("Valuation", level=1)
        low = val.get("valuation_range_low", {})
        mid = val.get("valuation_range_mid", {})
        high = val.get("valuation_range_high", {})
        doc.add_paragraph(
            f"Range: ${low.get('amount', 0):.1f}M - "
            f"${mid.get('amount', 0):.1f}M - "
            f"${high.get('amount', 0):.1f}M"
        )

    # Final Recommendation
    doc.add_heading("Final Recommendation", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"{memo_data.get('final_recommendation', 'N/A')}")
    run.bold = True
    run.font.size = Pt(14)
    p.add_run(f" (Conviction: {memo_data.get('conviction_level', 'N/A')})")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
