"""
Test script for multi-format document processing.

Creates sample PDF, DOCX, and XLSX files with pitch deck data
and tests the document service processing for each format.
"""

import asyncio
import os
import sys

# Add project to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from pathlib import Path


def create_sample_docx(output_path: str) -> None:
    """Create a sample DOCX pitch deck."""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Title
    title = doc.add_heading("TechFlow AI - Series A Pitch Deck", 0)

    # Company Overview
    doc.add_heading("Company Overview", level=1)
    doc.add_paragraph(
        "TechFlow AI is revolutionizing enterprise workflow automation using "
        "advanced artificial intelligence. Founded in 2022, we help Fortune 500 "
        "companies reduce operational costs by 40% through intelligent process automation."
    )

    # Key Metrics
    doc.add_heading("Key Metrics", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    metrics = [
        ("Annual Recurring Revenue (ARR)", "$8.5 Million"),
        ("YoY Growth Rate", "185%"),
        ("Gross Margin", "78%"),
        ("Monthly Burn Rate", "$450,000"),
        ("Runway", "18 months"),
        ("Total Customers", "127 enterprises"),
    ]

    for i, (metric, value) in enumerate(metrics):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value

    # Team
    doc.add_heading("Founding Team", level=1)
    doc.add_paragraph("Sarah Chen - CEO (ex-Google AI, Stanford CS PhD)")
    doc.add_paragraph("Michael Rodriguez - CTO (ex-Amazon, MIT)")
    doc.add_paragraph("Emily Watson - COO (ex-McKinsey, Harvard MBA)")

    # Market
    doc.add_heading("Market Opportunity", level=1)
    doc.add_paragraph(
        "Total Addressable Market (TAM): $85 Billion\n"
        "Serviceable Addressable Market (SAM): $25 Billion\n"
        "The enterprise automation market is growing at 23% CAGR."
    )

    # Funding
    doc.add_heading("Funding", level=1)
    doc.add_paragraph(
        "We are raising a $15 Million Series A at a $60 Million pre-money valuation.\n"
        "Previously raised: $3.5 Million (Seed from Sequoia, a16z)\n"
        "Use of funds: 50% R&D, 30% Sales & Marketing, 20% Operations"
    )

    # Competition
    doc.add_heading("Competition", level=1)
    doc.add_paragraph("Key competitors: UiPath ($2B+ raised), Automation Anywhere, Microsoft Power Automate")
    doc.add_paragraph("Our advantage: AI-first approach with 3x faster implementation time")

    doc.save(output_path)
    print(f"Created DOCX: {output_path}")


def create_sample_xlsx(output_path: str) -> None:
    """Create a sample XLSX financial model."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()

    # Summary Sheet
    ws = wb.active
    ws.title = "Company Summary"

    ws['A1'] = "TechFlow AI - Financial Summary"
    ws['A1'].font = Font(bold=True, size=14)

    summary_data = [
        ["", ""],
        ["Company Name", "TechFlow AI"],
        ["Industry", "Enterprise Software / AI"],
        ["Stage", "Series A"],
        ["Founded", "2022"],
        ["Headquarters", "San Francisco, CA"],
        ["Team Size", "45 employees"],
        ["", ""],
        ["KEY METRICS", ""],
        ["ARR", "$8,500,000"],
        ["MRR", "$708,333"],
        ["Growth Rate (YoY)", "185%"],
        ["Gross Margin", "78%"],
        ["Net Revenue Retention", "135%"],
        ["Churn Rate", "2.5%"],
        ["", ""],
        ["UNIT ECONOMICS", ""],
        ["CAC", "$12,500"],
        ["LTV", "$87,500"],
        ["LTV/CAC Ratio", "7.0x"],
        ["Payback Period", "8 months"],
        ["", ""],
        ["FUNDING", ""],
        ["Total Raised", "$3,500,000"],
        ["Current Round", "$15,000,000"],
        ["Pre-Money Valuation", "$60,000,000"],
        ["Burn Rate (Monthly)", "$450,000"],
        ["Runway", "18 months"],
    ]

    for i, (label, value) in enumerate(summary_data, start=3):
        ws[f'A{i}'] = label
        ws[f'B{i}'] = value
        if label and not value:
            ws[f'A{i}'].font = Font(bold=True)

    # Financial Projections Sheet
    ws2 = wb.create_sheet("Projections")

    headers = ["Year", "Revenue", "Growth %", "Gross Profit", "EBITDA", "Customers"]
    for col, header in enumerate(headers, start=1):
        ws2.cell(row=1, column=col, value=header).font = Font(bold=True)

    projections = [
        [2023, 3200000, 0, 2400000, -2500000, 45],
        [2024, 8500000, 166, 6630000, -1800000, 127],
        [2025, 22000000, 159, 17160000, 1500000, 310],
        [2026, 48000000, 118, 38400000, 8500000, 620],
        [2027, 85000000, 77, 68000000, 22000000, 1100],
    ]

    for row_idx, row_data in enumerate(projections, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            ws2.cell(row=row_idx, column=col_idx, value=value)

    # Market Sheet
    ws3 = wb.create_sheet("Market")

    ws3['A1'] = "Market Analysis"
    ws3['A1'].font = Font(bold=True, size=12)

    market_data = [
        ["", ""],
        ["TAM (Total Addressable Market)", "$85,000,000,000"],
        ["SAM (Serviceable Addressable Market)", "$25,000,000,000"],
        ["SOM (Serviceable Obtainable Market)", "$500,000,000"],
        ["Market Growth Rate", "23%"],
        ["", ""],
        ["Top Competitors", "Funding Raised"],
        ["UiPath", "$2,100,000,000"],
        ["Automation Anywhere", "$840,000,000"],
        ["Microsoft Power Automate", "N/A (Microsoft)"],
        ["Zapier", "$165,000,000"],
    ]

    for i, row in enumerate(market_data, start=2):
        ws3[f'A{i}'] = row[0]
        ws3[f'B{i}'] = row[1] if len(row) > 1 else ""

    # Founders Sheet
    ws4 = wb.create_sheet("Team")

    ws4['A1'] = "Founding Team"
    ws4['A1'].font = Font(bold=True, size=12)

    team_data = [
        ["Name", "Role", "Background", "Education"],
        ["Sarah Chen", "CEO", "Ex-Google AI Lead, 10 years ML experience", "Stanford CS PhD"],
        ["Michael Rodriguez", "CTO", "Ex-Amazon Principal Engineer", "MIT Computer Science"],
        ["Emily Watson", "COO", "Ex-McKinsey Partner", "Harvard MBA"],
    ]

    for i, row in enumerate(team_data, start=3):
        for j, value in enumerate(row):
            ws4.cell(row=i, column=j+1, value=value)
            if i == 3:
                ws4.cell(row=i, column=j+1).font = Font(bold=True)

    wb.save(output_path)
    print(f"Created XLSX: {output_path}")


def create_sample_pdf(output_path: str) -> None:
    """Create a simple text file as PDF placeholder (requires reportlab for real PDF)."""
    # For testing, we'll create a simple text content
    # In production, you'd use reportlab or similar
    content = """
    TECHFLOW AI - SERIES A PITCH DECK
    ==================================

    COMPANY OVERVIEW
    ----------------
    TechFlow AI is an enterprise workflow automation platform powered by AI.
    Founded: 2022
    Headquarters: San Francisco, CA
    Industry: Enterprise Software / Artificial Intelligence

    KEY FINANCIALS
    --------------
    Annual Recurring Revenue (ARR): $8.5 Million
    Year-over-Year Growth: 185%
    Gross Margin: 78%
    Monthly Burn Rate: $450,000
    Runway: 18 months

    TRACTION
    --------
    Total Customers: 127 enterprises
    Notable Customers: Fortune 500 companies
    Net Revenue Retention: 135%
    Churn Rate: 2.5%

    UNIT ECONOMICS
    --------------
    Customer Acquisition Cost (CAC): $12,500
    Lifetime Value (LTV): $87,500
    LTV/CAC Ratio: 7.0x
    Payback Period: 8 months

    MARKET OPPORTUNITY
    ------------------
    TAM: $85 Billion
    SAM: $25 Billion
    Market Growth Rate: 23% CAGR

    FOUNDING TEAM
    -------------
    Sarah Chen - CEO (ex-Google AI, Stanford PhD)
    Michael Rodriguez - CTO (ex-Amazon, MIT)
    Emily Watson - COO (ex-McKinsey, Harvard MBA)
    Team Size: 45 employees

    FUNDING
    -------
    Total Raised to Date: $3.5 Million (Seed)
    Current Round: $15 Million Series A
    Pre-Money Valuation: $60 Million
    Lead Investors: Sequoia, a16z

    USE OF FUNDS
    ------------
    50% - R&D and Product Development
    30% - Sales and Marketing
    20% - Operations and Hiring

    COMPETITION
    -----------
    - UiPath (raised $2.1B)
    - Automation Anywhere (raised $840M)
    - Microsoft Power Automate

    Our Advantage: AI-first approach with 3x faster implementation
    """

    # Try to create a real PDF if reportlab is available
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter

        # Write content line by line
        y = height - inch
        for line in content.strip().split('\n'):
            if y < inch:
                c.showPage()
                y = height - inch
            c.drawString(inch, y, line.strip())
            y -= 14

        c.save()
        print(f"Created PDF: {output_path}")
        return
    except ImportError:
        pass

    # Fallback: Create a text file with .txt extension for testing
    txt_path = output_path.replace('.pdf', '_content.txt')
    with open(txt_path, 'w') as f:
        f.write(content)
    print(f"Created text content (PDF library not available): {txt_path}")
    print("Note: Install reportlab for PDF creation: pip install reportlab")


async def test_document_processing():
    """Test document processing for all formats."""
    from app.services.document_service import DocumentService

    print("\n" + "="*60)
    print("MULTI-FORMAT DOCUMENT PROCESSING TEST")
    print("="*60 + "\n")

    # Create test directory
    test_dir = Path("test_documents")
    test_dir.mkdir(exist_ok=True)

    # Create sample documents
    print("Creating sample documents...\n")

    docx_path = str(test_dir / "techflow_pitch.docx")
    xlsx_path = str(test_dir / "techflow_financials.xlsx")
    pdf_path = str(test_dir / "techflow_deck.pdf")

    create_sample_docx(docx_path)
    create_sample_xlsx(xlsx_path)
    create_sample_pdf(pdf_path)

    print("\n" + "-"*60)
    print("Testing Document Service")
    print("-"*60 + "\n")

    # Initialize service
    service = DocumentService()

    # Test each format
    test_files = [
        ("DOCX", docx_path),
        ("XLSX", xlsx_path),
    ]

    # Only test PDF if it was created (requires reportlab)
    if Path(pdf_path).exists():
        test_files.append(("PDF", pdf_path))

    results = {}

    for format_name, file_path in test_files:
        print(f"\n{'='*40}")
        print(f"Testing {format_name}: {Path(file_path).name}")
        print(f"{'='*40}")

        try:
            images, text = await service.process_document(file_path)

            print(f"  Images extracted: {len(images)}")
            print(f"  Text length: {len(text)} characters")
            print(f"  Text preview (first 500 chars):")
            print("-"*40)
            print(text[:500] if text else "(no text)")
            print("-"*40)

            results[format_name] = {
                "success": True,
                "images": len(images),
                "text_length": len(text),
            }

        except Exception as e:
            print(f"  ERROR: {type(e).__name__}: {e}")
            results[format_name] = {
                "success": False,
                "error": str(e),
            }

    # Summary
    print("\n" + "="*60)
    print("TEST RESULTS SUMMARY")
    print("="*60)

    for format_name, result in results.items():
        status = "PASS" if result.get("success") else "FAIL"
        print(f"\n{format_name}: {status}")
        if result.get("success"):
            print(f"  - Images: {result['images']}")
            print(f"  - Text: {result['text_length']} chars")
        else:
            print(f"  - Error: {result.get('error')}")

    print("\n" + "="*60)
    return results


async def test_extraction_with_cache():
    """Test the extraction agent with caching."""
    from app.services.document_service import DocumentService
    from app.agents.extraction_agent import ExtractionAgent
    from app.services.cache_service import get_extraction_cache
    import time

    print("\n" + "="*60)
    print("EXTRACTION AGENT TEST WITH CACHING")
    print("="*60 + "\n")

    test_dir = Path("test_documents")
    docx_path = str(test_dir / "techflow_pitch.docx")

    if not Path(docx_path).exists():
        print("Test documents not found. Run document processing test first.")
        return

    # Process document
    service = DocumentService()
    images, text = await service.process_document(docx_path)

    print(f"Document processed: {len(text)} chars of text\n")

    # Initialize extraction agent
    agent = ExtractionAgent()
    cache = get_extraction_cache()

    # Clear cache for clean test
    cache.clear()
    print("Cache cleared for testing\n")

    # First extraction (should hit API)
    print("First extraction (API call expected)...")
    start = time.time()
    try:
        result1 = await agent.execute(
            images=images,
            text_content=text,
            extraction_mode="quick"  # Use quick mode for faster test
        )
        elapsed1 = time.time() - start
        print(f"  Completed in {elapsed1:.2f}s")
        print(f"  Company: {result1.company_name}")
        print(f"  Confidence: {result1.extraction_confidence:.2f}")
    except Exception as e:
        print(f"  ERROR: {e}")
        return

    # Second extraction (should hit cache)
    print("\nSecond extraction (cache hit expected)...")
    start = time.time()
    try:
        result2 = await agent.execute(
            images=images,
            text_content=text,
            extraction_mode="quick"
        )
        elapsed2 = time.time() - start
        print(f"  Completed in {elapsed2:.2f}s")
        print(f"  Company: {result2.company_name}")

        if elapsed2 < elapsed1 / 2:
            print(f"\n  CACHE WORKING! Second call was {elapsed1/elapsed2:.1f}x faster")

    except Exception as e:
        print(f"  ERROR: {e}")

    # Show cache stats
    stats = cache.get_stats()
    print(f"\nCache stats: {stats}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Test multi-format document processing")
    parser.add_argument("--extraction", action="store_true", help="Also test extraction agent")
    args = parser.parse_args()

    # Run document processing test
    asyncio.run(test_document_processing())

    # Optionally test extraction
    if args.extraction:
        asyncio.run(test_extraction_with_cache())
